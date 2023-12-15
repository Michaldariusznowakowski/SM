from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO
import cv2
from tqdm import tqdm
import sys


class DocxSave:
    def __init__(self, debug=False):
        self.document = Document()
        self.debug = debug

    def h(self, text, level):
        if self.debug:
            print("h: "+text)
        else:
            self.document.add_heading(text, level)

    def p(self, text):
        if self.debug:
            print("p: "+text)
        else:
            self.document.add_paragraph(text)

    def addImage(self, image, width):
        if self.debug:
            plt.show()
        else:
            if isinstance(image, plt.Figure):
                memfile = BytesIO()
                image.savefig(memfile)
                self.document.add_picture(memfile, width=Inches(width))
                memfile.close()
                plt.close()

    def save(self, filename):
        if not self.debug:
            self.document.save(filename)


def RLE(data: np.ndarray) -> np.ndarray:
    input = data.astype(np.int32)
    Array1D = input.flatten()
    out = []
    counter = 1
    for i in tqdm(range(1, len(Array1D))):
        if (Array1D[i] == Array1D[i-1]):
            counter += 1
        else:
            out.append(counter)
            out.append(Array1D[i-1])
            counter = 1
    out.append(counter)
    out.append(Array1D[i])
    output = np.array([len(input.shape)])
    output = np.concatenate([output, input.shape])
    output = np.concatenate([output, out])
    return output


def RLEDecode(data: np.ndarray) -> np.ndarray:
    nd = data[0]
    shape = tuple(data[1:nd+1])
    raw = data[nd+1:]
    out = []
    for i in tqdm(range(0, len(raw), 2)):
        for j in range(raw[i]):
            out.append(raw[i+1])
    return np.array(out).reshape(shape)


def get_size(obj):
    return len(obj)


def get_size_rle(obj):
    nd = obj[0]
    shape = tuple(obj[1:nd+1])
    raw = obj[nd+1:]
    return len(raw)


def divideIntoBlocks(img):  # 128x128
    width = img.shape[1]
    height = img.shape[0]
    blocks = []
    for w in range(0, width, 128):
        for k in range(0, height, 128):
            blocks.append(img[k:(k+128), w:(w+128)])
    return blocks


def getRandomBlock(blocks, count=1):
    output = []
    found = 0
    while found < count:
        idx = np.random.randint(0, len(blocks))
        if idx not in output:
            if blocks[idx].shape[0] == 128 and blocks[idx].shape[1] == 128:
                output.append(idx)
                found += 1
    return [blocks[i] for i in output]


####################################################################################################
# Lab 7
####################################################################################################
QY = np.array([
    [16, 11, 10, 16, 24,  40,  51,  61],
    [12, 12, 14, 19, 26,  58,  60,  55],
    [14, 13, 16, 24, 40,  57,  69,  56],
    [14, 17, 22, 29, 51,  87,  80,  62],
    [18, 22, 37, 56, 68,  109, 103, 77],
    [24, 36, 55, 64, 81,  104, 113, 92],
    [49, 64, 78, 87, 103, 121, 120, 101],
    [72, 92, 95, 98, 112, 100, 103, 99],
])
QC = np.array([
    [17, 18, 24, 47, 99, 99, 99, 99],
    [18, 21, 26, 66, 99, 99, 99, 99],
    [24, 26, 56, 99, 99, 99, 99, 99],
    [47, 66, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99],
])


def toYCrCB(rgb):
    return cv2.cvtColor(rgb, cv2.COLOR_RGB2YCrCb).astype(int)


def toRGB(ycrcb):
    return cv2.cvtColor(ycrcb.astype(np.uint8), cv2.COLOR_YCrCb2RGB)


def DCT(img):
    return cv2.dct(img.astype(np.float32))


def IDCT(img):
    return cv2.idct(img.astype(np.float32))


class DataContainer:
    def __init__(self, Y, Cb, Cr, sh, Ratio, QY=QY, QC=QC):
        self.shape = sh
        self.Y = Y
        self.Cb = Cb
        self.Cr = Cr
        self.ChromaRatio = Ratio
        self.QY = QY
        self.QC = QC


def verX(Ratio):
    if Ratio == "4:2:2":
        return JPEG422()
    else:
        return JPEG444()


class JPEG422:
    def __init__(self, con: DataContainer = None):
        if con is not None:
            self.data_container = con
        else:
            self.data_container = DataContainer(np.array([]), np.array(
                []), np.array([]), (0, 0), "4:2:2")

    def __call__(self, YCrCb):
        self.data_container.shape = YCrCb.shape
        self.data_container.Y = YCrCb[:, :, 0]
        self.data_container.Cr = YCrCb[:, :, 1]
        self.data_container.Cb = YCrCb[:, :, 2]
        return self

    def ChromaSubsampling(self):
        self.data_container.Cb[:, 1::2] = self.data_container.Cb[:, ::2]
        self.data_container.Cr[:, 1::2] = self.data_container.Cr[:, ::2]
        return self

    def ChromaResampling(self):
        self.data_container.Cb = np.repeat(self.data_container.Cb, 2, axis=1)
        self.data_container.Cr = np.repeat(self.data_container.Cr, 2, axis=1)
        return self

    def update(self, Y, Cb, Cr):
        self.data_container.Y = Y
        self.data_container.Cr = Cr
        self.data_container.Cb = Cb
        return self


class JPEG444:
    def __init__(self, data_container=None):
        if data_container is not None:
            self.data_container = data_container
        else:
            self.data_container = DataContainer(np.array([]), np.array(
                []), np.array([]), (0, 0), "4:4:4")

    def __call__(self, YCrCb):
        self.data_container.shape = YCrCb.shape
        self.data_container.Y = YCrCb[:, :, 0]
        self.data_container.Cr = YCrCb[:, :, 1]
        self.data_container.Cb = YCrCb[:, :, 2]
        return self

    def update(self, Y, Cb, Cr):
        self.data_container.Y = Y
        self.data_container.Cr = Cr
        self.data_container.Cb = Cb
        return self

    def ChromaSubsampling(self):
        return self

    def ChromaResampling(self):
        return self


def ZigZag(A):
    template = np.array([
        [0,  1,  5,  6,  14, 15, 27, 28],
        [2,  4,  7,  13, 16, 26, 29, 42],
        [3,  8,  12, 17, 25, 30, 41, 43],
        [9,  11, 18, 24, 31, 40, 44, 53],
        [10, 19, 23, 32, 39, 45, 52, 54],
        [20, 22, 33, 38, 46, 51, 55, 60],
        [21, 34, 37, 47, 50, 56, 59, 61],
        [35, 36, 48, 49, 57, 58, 62, 63],
    ])
    if len(A.shape) == 1:
        B = np.zeros((8, 8))
        for r in range(0, 8):
            for c in range(0, 8):
                B[r, c] = A[template[r, c]]
    else:
        B = np.zeros((64,))
        for r in range(0, 8):
            for c in range(0, 8):
                B[template[r, c]] = A[r, c]
    return B


def CompressBlock(block, Q):
    output = DCT(block - 128)
    output = np.round(output / Q).astype(int)
    output = ZigZag(output)
    return output


def DecompressBlock(vector, Q):
    output = ZigZag(vector)
    output = output * Q
    output = IDCT(output) + 128
    return output


def DecompressLayer(S, Q):
    L = np.zeros((128, 128))
    for idx, i in enumerate(range(0, S.shape[0], 64)):
        vector = S[i:(i+64)]
        m = L.shape[0]/8
        k = int((idx % m)*8)
        w = int((idx//m)*8)
        L[w:(w+8), k:(k+8)] = DecompressBlock(vector, Q)
    return L


def CompressLayer(L, Q):
    S = np.array([])
    for w in range(0, L.shape[0], 8):
        for k in range(0, L.shape[1], 8):
            block = L[w:(w+8), k:(k+8)]
            S = np.append(S, CompressBlock(block, Q))
    return S


def CompressJPEG(RGB, Ratio="4:4:4", QY=QY, QC=QC):
    YCrCb = toYCrCB(RGB)
    JPEG = verX(Ratio)
    JPEG(YCrCb)
    JPEG.data_container.QC = QC
    JPEG.data_container.QY = QY
    JPEG.ChromaSubsampling()
    JPEG.data_container.Y = CompressLayer(
        JPEG.data_container.Y, JPEG.data_container.QY)
    JPEG.data_container.Cr = CompressLayer(
        JPEG.data_container.Cr, JPEG.data_container.QC)
    JPEG.data_container.Cb = CompressLayer(
        JPEG.data_container.Cb, JPEG.data_container.QC)
    size = len(JPEG.data_container.Y)+len(
        JPEG.data_container.Cr)+len(JPEG.data_container.Cb)
    JPEG.data_container.Y = RLE(JPEG.data_container.Y)
    JPEG.data_container.Cr = RLE(JPEG.data_container.Cr)
    JPEG.data_container.Cb = RLE(JPEG.data_container.Cb)
    return JPEG.data_container, size


def DecompressJPEG(data_container):
    JPEG = verX(data_container.ChromaRatio)
    JPEG.data_container = data_container
    JPEG.data_container.Y = RLEDecode(JPEG.data_container.Y)
    JPEG.data_container.Cr = RLEDecode(JPEG.data_container.Cr)
    JPEG.data_container.Cb = RLEDecode(JPEG.data_container.Cb)
    Y = DecompressLayer(JPEG.data_container.Y, JPEG.data_container.QY)
    Cr = DecompressLayer(
        JPEG.data_container.Cr, JPEG.data_container.QC)
    Cb = DecompressLayer(
        JPEG.data_container.Cb, JPEG.data_container.QC)
    JPEG.update(Y, Cb, Cr)
    JPEG.ChromaResampling()
    Cr = JPEG.data_container.Cr
    Cb = JPEG.data_container.Cb
    Y = cv2.resize(Y, (Cr.shape[1], Cr.shape[0]),
                   interpolation=cv2.INTER_LINEAR)
    YCrCb = cv2.resize(np.dstack([Y, Cr, Cb]).clip(0, 255).astype(np.uint8), (JPEG.data_container.shape[1], JPEG.data_container.shape[0]),
                       interpolation=cv2.INTER_LINEAR)
    return toRGB(YCrCb)


def start(Debug=False, images_folder="data/", image_name="zut.jpg"):
    dx = DocxSave(debug=Debug)
    dx.h("Implementacja częściowej kompresji JPEG", 1)
    dx.h("Michał Nowakowski nm49353", 2)
    dx.p("Użyty algorytm: RLE")
    img = cv2.imread(images_folder+image_name)
    fig, axs = plt.subplots(1, 1)
    axs.axis('off')
    axs.imshow(img)
    dx.p("Obraz wejściowy")
    dx.addImage(fig, 6)
    dx.p("Losowe bloki")
    allblocks = divideIntoBlocks(img)
    blocks = getRandomBlock(allblocks, 4)
    counted = len(blocks)
    fig, axs = plt.subplots(counted, 1)
    for idx, block in enumerate(blocks):
        axs[idx].imshow(block)
        axs[idx].axis('off')
    dx.addImage(fig, 6)
    for img, img_index in zip(blocks, range(len(blocks))):
        img_before_comp = img
        ratios = ["4:4:4", "4:2:2"]
        for ratio in ratios:
            dx.h("Kompresja JPEG "+ratio, 2)
            before_y = img[:, :, 0]
            before_Cr = img[:, :, 1]
            before_Cb = img[:, :, 2]
            con, size_before_rle = CompressJPEG(img, ratio)
            size_after_rle = get_size_rle(
                con.Y)+get_size_rle(con.Cr)+get_size_rle(con.Cb)
            after_img = DecompressJPEG(con)
            dx.p("Rozmiar po rle: "+str(size_after_rle))
            dx.p("Rozmiar przed rle: " +
                 str(size_before_rle))
            dx.p("Kompresja: "+str(size_after_rle /
                                   size_before_rle*100)+"%")
            ycrb = toYCrCB(after_img)
            after_y = ycrb[:, :, 0]
            after_Cr = ycrb[:, :, 1]
            after_Cb = ycrb[:, :, 2]
            fig, axs = plt.subplots(4, 2, sharey=True)
            fig.set_size_inches(9, 13)
            axs[0, 0].imshow(img_before_comp)
            axs[1, 0].imshow(before_y, cmap=plt.cm.gray)
            axs[2, 0].imshow(before_Cr, cmap=plt.cm.gray)
            axs[3, 0].imshow(before_Cb, cmap=plt.cm.gray)
            axs[0, 1].imshow(after_img)
            axs[1, 1].imshow(after_y, cmap=plt.cm.gray)
            axs[2, 1].imshow(after_Cr, cmap=plt.cm.gray)
            axs[3, 1].imshow(after_Cb, cmap=plt.cm.gray)
            for row in axs:
                for ax in row:
                    ax.axis('off')
            dx.addImage(fig, 5)
            cv2.imwrite("data/after_"+image_name+"_"+ratio.replace(":", "") +
                        "_"+str(img_index)+".jpg", after_img)
            cv2.imwrite("data/before_"+image_name+"_"+ratio.replace(":", "") +
                        "_"+str(img_index)+".jpg", img_before_comp)
            dx.save("lab_7_nm49353_michal_nowakowski.docx")


if "__main__" == __name__:
    DEBUG = False  # Brak zapisu do pliku, wyświetlanie wykresów
    images_folder = "data/"
    image_name = "zut.jpg"
    start(Debug=DEBUG, images_folder=images_folder, image_name=image_name)
    ###############################################################

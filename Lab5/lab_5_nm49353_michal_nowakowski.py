from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO
from tqdm import tqdm
import sys
####################################################################################################
# Lab 5
####################################################################################################

DEBUG = False # Brak zapisu do pliku, wyświetlanie wykresów
DISABLE_PLOTS = False # Wyłączenie wyświetlania wykresów
dataFolder = "data/"

####################################################################################################
class DocxSave:
    def __init__(self, debug=False):
        self.document = Document()
        self.debug = debug
        
    def h(self, text, level):
        if self.debug:
            print("h: "+text)
        else:
            self.document.add_heading(text, level)
    
    def p(self, text):
        if self.debug:
            print("p: "+text)
        else:
            self.document.add_paragraph(text)
    
    def addImage(self, image, width):
        if self.debug and DISABLE_PLOTS == False:
            plt.show()
        else:
            if isinstance(image, plt.Figure):
                memfile = BytesIO()
                image.savefig(memfile)
                self.document.add_picture(memfile, width=Inches(width))
                memfile.close()
                plt.close()
    
    def addCode(self, codePath):
        if not self.debug: 
            with open(codePath, 'r') as codeFile:
                codeText = codeFile.read()
                self.document.add_paragraph('Kod źródłowy:')
                codeParagraph = self.document.add_paragraph()
                codeRun = codeParagraph.add_run(codeText)
                codeRun.font.name = 'Courier New'
    
    def save(self, filename):
        if not self.debug:
            self.document.save(filename)
            
dx = DocxSave(debug=DEBUG)



def get_size(obj, seen=None):
    """Recursively finds size of objects"""
    size = sys.getsizeof(obj)
    if seen is None:
        seen = set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    # Important mark as seen *before* entering recursion to gracefully handle
    # self-referential objects
    seen.add(obj_id)
    if isinstance(obj,np.ndarray):
        size=obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v, seen) for v in obj.values()])
        size += sum([get_size(k, seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size
####################################################################################################

def RLE(data: np.ndarray) -> np.ndarray:
    input=data.astype(np.int32)
    Array1D=input.flatten()
    out=[]
    counter=1
    for i in tqdm(range(1,len(Array1D))):
        if (Array1D[i] == Array1D[i-1]):
            counter+=1
        else:
            out.append(counter)
            out.append(Array1D[i-1])
            counter=1
    out.append(counter)
    out.append(Array1D[i])
    output=np.array([len(input.shape)])
    output=np.concatenate([output,input.shape])
    output=np.concatenate([output,out])
    return output

def RLEDecode(data: np.ndarray) -> np.ndarray:
    nd=data[0]
    shape=tuple(data[1:nd+1])
    raw=data[nd+1:]
    out=[]
    for i in tqdm(range(0,len(raw),2)):
        for j in range(raw[i]):
            out.append(raw[i+1])
    return np.array(out).reshape(shape)

def ByteRun(data: np.ndarray) -> np.ndarray:
    input=data.astype(np.int32)
    Array1D=input.flatten()
    index=0
    new=True
    i = 0
    out = []
    with tqdm(total=len(Array1D)) as pbar:
        while i < len(Array1D):
            if (new):
                if (i != len(Array1D)-1 and Array1D[i] == Array1D[i+1]):
                    out.append(0)
                    out.append(Array1D[i])
                    index=len(out)-2
                    new=False
                else:
                    out.append(Array1D[i])
                    new=True 
            else:
                if (out[index] == 126):
                    New = True
                    continue
                if (Array1D[i] == out[index+1]):
                    out[index] -= 1
                else:
                    new=True
                    continue
            i+=1
            pbar.update(1)
                
    output = np.array([len(input.shape)])
    output = np.concatenate([output, input.shape])
    output = np.concatenate([output, out])
    return output

def ByteRunDecode(data: np.ndarray) -> np.ndarray:
    nd=data[0]
    shape=tuple(data[1:nd+1])
    raw=data[nd+1:]
    out=[]
    for i in tqdm(range(0,len(raw))):
        if (raw[i]<0):
            for j in range(-raw[i]):
                out.append(raw[i+1])
        else:
            out.append(raw[i])
    return np.array(out).reshape(shape)
            
files = [ dataFolder+"rdoc.jpg", dataFolder+"rtech.jpg",dataFolder+"rcolor.jpg"]

dx.h("Zadanie 1", 1)

tests = [np.array([1,1,1,1,2,1,1,1,1,2,1,1,1,1]),np.array([1,2,3,1,2,3,1,2,3]),np.array([5,1,5,1,5,5,1,1,5,5,1,1,5]),np.array([-1,-1,-1,-5,-5,-3,-4,-2,1,2,2,1]),np.zeros((1,520)),np.arange(0,521,1),np.eye(7),np.dstack([np.eye(7),np.eye(7),np.eye(7)]),np.ones((1,1,1,1,1,1,10))]
dx.h("Zadanie 1.1 - RLE", 2)  
for test in tests:
    rle = RLE(test)
    rleDec = RLEDecode(rle)
    flatIn = test.flatten()
    rleND = rle[0]
    rleShape = tuple(rle[1:rleND+1])
    rleFlatOut = rle[rleND+1:]
    dx.h("Test: "+str(test),3)
    dx.p("Rle wejscie: "+str(flatIn))
    dx.p("Rle po zakodowaniu: "+str(rleFlatOut))
    dx.p("Rle po dekodowaniu: "+str(rleDec))
    dx.p("Liczba elementów na wejściu: "+str(len(flatIn)))
    dx.p("Liczba elementów na wyjściu: "+str(len(rleFlatOut)))
    dx.p("Stopień kompresji: "+str(len(flatIn)/len(rleFlatOut)))
    dx.p("Procent kompresji: "+str(len(rleFlatOut)/len(flatIn)*100)+"%")

tests = [np.array([1,1,1,1,2,1,1,1,1,2,1,1,1,1]),np.array([1,2,3,1,2,3,1,2,3]),np.array([5,1,5,1,5,5,1,1,5,5,1,1,5]),np.zeros((1,520)),np.arange(0,521,1),np.eye(7),np.dstack([np.eye(7),np.eye(7),np.eye(7)]),np.ones((1,1,1,1,1,1,10))]
dx.h("Zadanie 1.2 - ByteRun", 2)
for test in tests:
    br = ByteRun(test)
    brDec = ByteRunDecode(br)
    
    flatIn = test.flatten()
    
    brND = br[0]
    brShape = tuple(br[1:brND+1])
    brFlatOut = br[brND+1:]
    
    dx.h("Test: "+str(test),3)
    dx.p("ByteRun wejscie: "+str(flatIn))
    dx.p("ByteRun po zakodowaniu: "+str(brFlatOut))
    dx.p("ByteRun po dekodowaniu: "+str(brDec))
    dx.p("Liczba elementów na wejściu: "+str(len(flatIn)))
    dx.p("Liczba elementów na wyjściu: "+str(len(brFlatOut)))
    dx.p("Stopień kompresji: "+str(len(flatIn)/len(brFlatOut)))
    dx.p("Procent kompresji: "+str(len(brFlatOut)/len(flatIn)*100)+"%")

dx.h ("Zadanie 2", 1)
for filename in files:
    image = plt.imread(filename)
    br=ByteRun(image)
    brND=br[0]
    brShape=tuple(br[1:brND+1])
    brFlatOut=br[brND+1:]
    
    rle=RLE(image)
    rleND=rle[0]
    rleShape=tuple(rle[1:rleND+1])
    rleFlatOut=rle[rleND+1:]
    
    rleDec=RLEDecode(rle)
    brDec=ByteRunDecode(br)
    
    dx.h("Nazwa pliku: "+filename,2)
    fig, ax  = plt.subplots(1, 3, figsize=(10, 5))
    ax[0].imshow((image.astype(np.float32) / 255.0))
    ax[0].set_title("Obraz oryginalny")
    ax[1].imshow((rleDec.astype(np.float32) / 255.0))
    ax[1].set_title("Obraz po RLE")
    ax[2].imshow((brDec.astype(np.float32) / 255.0))
    ax[2].set_title("Obraz po ByteRun")
    dx.p("Rozmiar obrazu: "+str(get_size(image))+" bajtów")
    dx.p("Rozmiar po RLE: "+str(get_size(brFlatOut))+" bajtów")
    dx.p("Rozmiar po ByteRun: "+str(get_size(rleFlatOut))+" bajtów")
    dx.p("Stopień kompresji RLE: "+str(get_size(image)/get_size(brFlatOut)))
    dx.p("Stopień kompresji ByteRun: "+str(get_size(image)/get_size(rleFlatOut)))
    dx.p("Procent kompresji RLE: "+str(get_size(brFlatOut)/get_size(image)*100)+"%")
    dx.p("Procent kompresji ByteRun: "+str(get_size(rleFlatOut)/get_size(image)*100)+"%")
    dx.addImage(fig, 5)

dx.h("Wnioski",1)
dx.h("rdoc.jpg",2)
dx.p("Orginalny obraz jest jednolity i nie zawiera szumów, o wiele lepiej sprawdza się byteRun.")

dx.h("rtech.jpg",2)
dx.p("Orginalny obraz nie jest jednolity i zawiera szumy, wyniki są bardzo zbliżone, ale nieco lepiej wypada ByteRun.")

dx.h("rcolor.jpg",2)
dx.p("Dla tego obrazu, najlepiej sprawdza się ByteRun, ponieważ obraz zawiera różnych kolorów, co powoduje, że RLE nie jest w stanie skompresować go tak dobrze jak ByteRun. ")

dx.h("Wniosek ogólny",2)
dx.p("Pomimo prostoty działania tych algorytmów świetnie nadają się do kompresji, znacząco zmniejszając rozmiar plików.")
dx.p("A dodatkowo nie wpływają na jakość obrazu, ponieważ są to algorytmy bezstratne.")
print("Done")

dx.save("lab_5_nm49353_michal_nowakowski.docx")
    

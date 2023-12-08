from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO
import sounddevice as sd
import soundfile as sf
from scipy.interpolate import interp1d
from scipy import signal
import scipy.fftpack
import os.path
####################################################################################################
# Lab 6
####################################################################################################

DEBUG = True # Brak zapisu do pliku, wyświetlanie wykresów
DISABLE_PLOTS = False # Wyłączenie wyświetlania wykresów
dataFolder = "data/"

    
####################################################################################################
class DocxSave:
    def __init__(self, debug=False):
        self.document = Document()
        self.debug = debug
        
    def h(self, text, level):
        if self.debug:
            print("h: "+text)
        else:
            self.document.add_heading(text, level)
    
    def p(self, text):
        if self.debug:
            print("p: "+text)
        else:
            self.document.add_paragraph(text)
    
    def addImage(self, image, width):
        if self.debug and DISABLE_PLOTS == False:
            plt.show()
        else:
            if isinstance(image, plt.Figure):
                memfile = BytesIO()
                image.savefig(memfile)
                self.document.add_picture(memfile, width=Inches(width))
                memfile.close()
                plt.close()
    
    def addCode(self, codePath):
        if not self.debug: 
            with open(codePath, 'r') as codeFile:
                codeText = codeFile.read()
                self.document.add_paragraph('Kod źródłowy:')
                codeParagraph = self.document.add_paragraph()
                codeRun = codeParagraph.add_run(codeText)
                codeRun.font.name = 'Courier New'
    
    def save(self, filename):
        if not self.debug:
            self.document.save(filename)

def plotAudio(Signal, Fs, TimeMargin=[0, 0.04]):
    start = int(TimeMargin[0] * Fs)
    end = int(TimeMargin[1] * Fs)
    signalFragment = Signal[start:end]
    # time in sec
    time = np.linspace(0, len(signalFragment) / Fs, num=len(signalFragment))
    fig, axs = plt.subplots(2, 1, figsize=(10, 15))
    axs[0].plot(time, signalFragment)
    axs[0].set_title('Sygnał dźwiękowy')
    axs[0].set_xlabel('Czas [s]')
    axs[0].set_ylabel('Amplituda')
    spectrum = 20 * np.log10(np.abs(np.fft.fft(signalFragment))+np.finfo(np.float32).eps)
    freq = np.fft.fftfreq(len(signalFragment), 1 / Fs)
    halfSpectrum = spectrum[:len(spectrum)//2]
    halfFreq = freq[:len(freq)//2]
    axs[1].plot(halfFreq, halfSpectrum)
    axs[1].set_title('Widmo dźwiękowe (połowa)')
    axs[1].set_xlabel('Częstotliwość [Hz]')
    axs[1].set_ylabel('Amplituda [dB]')
    fig.tight_layout()
    return fig

def saveAudio(Signal, Fs, filename):
    file = dataFolder+"output/"+filename.replace(".wav", "")+".wav"
    if os.path.isfile(file):
        os.remove(file)
                
    sf.write(file, Signal, Fs)

dx = DocxSave(debug=DEBUG)
####################################################################################################

def Kwant(data,bit):
    
    output = []
    i = 0
    for x in data:
        if x >= 0:
            output.append(int(x*(2**bit-1)))
        else:
            output.append(int(x*(2**bit)))
        i+=1
    return output

    
        
def Decymacja(data,fs,decymacja):
    output = data[::decymacja]
    newFs = fs//decymacja
    return output, newFs

def Interpolacja(data,fs,fs2,typ):
    duration = data.shape[0]/fs
    x_old = np.arange(0, duration, 1/fs)
    x_new = np.arange(0, duration, 1/fs2)
    interp_func = interp1d(x_old, data, typ, fill_value="extrapolate")
    return interp_func(x_new).astype(data.dtype),fs2
####################################################################################################


def muLaw(signal, mu = 255):
    signal = signal.astype(np.float32)
    output = []
    for x in signal:
        output.append(np.sign(x) * np.log(1 + mu * np.abs(x)) / np.log(1 + mu))
    return output


def muLawDecode(signal, mu = 255):
    output = []
    for x in signal:
        output.append(np.sign(x) * (1 / mu) * (np.power(1 + mu, np.abs(x)) - 1))
    return output


def aLaw(signal, a = 87.6):
    signal = signal.astype(np.float32)
    output = []
    for x in signal:
        if np.abs(x) < 1/a:
            output.append(a * np.abs(x) / (1 + np.log(a)))
        else:
            output.append((1 + np.log(a * np.abs(x))) / (1 + np.log(a)))
    return np.array(output)

def aLawDecode(signal, a = 87.6):
    output = []
    for x in signal:
        if np.abs(x) < 1/np.log(a):
            output.append(np.sign(x)*(np.abs(x)*(1+np.log(a))/a))
        else:
            output.append(np.sign(x)*(np.exp(np.abs(x)*(1+np.log(a))-1))/a)
    return np.array(output)


dx.h("Lab 6", 0)
dx.p("Michał Nowakowski, Numer indeksu: 49353")
dx.h("Zadanie 1 Na zajęciach", 1)
x=np.linspace(-1,1,1000)
y=0.9*np.sin(np.pi*x*4)


A = 87.6
bits = 8
al_y=aLaw(y,A)
al_y8=Kwant(al_y,bits)
al_x=aLawDecode(al_y,A)

mu=255
bits = 8
mul_y=muLaw(y)
mul_y8=Kwant(mul_y,bits)
mul_x=muLawDecode(mul_y,mu)


dx.h("ALaw", 2)
dx.p("A = "+str(A))
dx.p("bits = "+str(bits))
dx.p("y = "+str(y))
dx.p("al_y = "+str(al_y))
dx.p("al_y8 = "+str(al_y8))
dx.p("al_x = "+str(al_x))

dx.h("MuLaw", 2)
dx.p("mu = "+str(mu))
dx.p("bits = "+str(bits))
dx.p("y = "+str(y))
dx.p("mul_y = "+str(mul_y))
dx.p("mul_y8 = "+str(mul_y8))
dx.p("mul_x = "+str(mul_x))

fig, axs = plt.subplots(1, 2, figsize=(20, 20))
# sygnal po kompresji a-law po kwantyzacji 8 bitowej
axs[0].set_xlabel('Wartość sygnału wejściowego')
axs[0].set_ylabel('Wartość sygnału wyjściowego')
axs[0].plot(y,al_y8,'r', label='Sygnal po kompresji a-law po kwantyzacji 8 bitowej')
axs[0].plot(y,y,'g', label='Sygnal po kompresji a-law bez kwantyzacji')
axs[0].plot(y,mul_y8,'b', label='Sygnal po kompresji mu-law po kwantyzacji 8 bitowej')
axs[0].plot(y,mul_y,'p', label='Sygnal po kompresji mu-law bez kwantyzacji')
axs[0].legend()
axs[0].grid(True)
axs[0].set_title("Krzywa kompresji")

# sygnal po dekompresji a-law po kwantyzacji 8 bitowej
axs[1].set_xlabel('Wartość sygnału wejściowego')
axs[1].set_ylabel('Wartość sygnału wyjściowego')
axs[1].plot(y,al_x,'r', label='Sygnal po dekompresji a-law po kwantyzacji 8 bitowej')
axs[1].plot(y,mul_x,'g', label='Sygnal po dekompresji mu-law po kwantyzacji 8 bitowej')
axs[1].plot(y,Kwant(y,8),'b', label='Sygnal po kwantyzacji 8 bitowej')
axs[1].legend()
axs[1].grid(True)
axs[1].set_title("Krzywa dekompresji")
dx.addImage(fig,4)
    
dx.save(dataFolder+"output/Lab_6.docx")
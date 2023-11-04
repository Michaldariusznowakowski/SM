from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO

class DocxSave:
    def __init__(self, debug=False):
        self.document = Document()
        self.debug = debug
        
    def h(self, text, level):
        if self.debug:
            print("h: "+text)
        else:
            self.document.add_heading(text, level)
    
    def p(self, text):
        if self.debug:
            print("p: "+text)
        else:
            self.document.add_paragraph(text)
    
    def addImage(self, image, width):
        if self.debug:
            plt.show()
        else:
            if isinstance(image, plt.Figure):
                memfile = BytesIO()
                image.savefig(memfile)
                self.document.add_picture(memfile, width=Inches(width))
                memfile.close()
                plt.close()
    
    def addCode(self, codePath):
        if not self.debug: 
            with open(codePath, 'r') as codeFile:
                codeText = codeFile.read()
                self.document.add_paragraph('Kod źródłowy:')
                codeParagraph = self.document.add_paragraph()
                codeRun = codeParagraph.add_run(codeText)
                codeRun.font.name = 'Courier New'
    
    def save(self, filename):
        if not self.debug:
            self.document.save(filename)


####################################################################################################
# Lab 3
####################################################################################################

DEBUG = False # Brak zapisu do pliku, wyświetlanie wykresów
dx = DocxSave(debug=DEBUG)
images_folder = "data/"

####################################################################################################


def colorFit(pixel,Pallete):
    return Pallete[np.argmin(np.linalg.norm(Pallete-pixel,axis=1))]

N = 3
paleta = np.linspace(0,1,N).reshape(N,1)

dx.h("Lab 3",0)
dx.h("Kwantyzacja przy użyciu dopasowywania do wzorca lub poszukiwanie najbliższych wartości koloru w palecie",1)
dx.p("Dla N = "+str(N))
dx.p("Paleta kolorów")
dx.p(str(paleta))


dx.p("Test funkcji colorFit")
paleta = np.linspace(0,1,3).reshape(3,1)
dx.p("Paleta kolorów")
dx.p("np.linspace(0,1,3).reshape(3,1)")
dx.p(str(paleta))
dx.p("colorFit(0.43,paleta)")
dx.p(str(colorFit(0.43,paleta)))
dx.p("colorFit(0.66,paleta)")
dx.p(str(colorFit(0.66,paleta)))
dx.p("colorFit(0.8,paleta)")
dx.p(str(colorFit(0.8,paleta)))

dx.h("Przedefiniowane palety kolorów",1)
pallet1 = np.array([[0], [1]])
pallet2 = np.array([[0], [0.5], [1]])
pallet4 = np.array([[0], [0.33], [0.66], [1]])

pallet8 = np.array([
        [0.0, 0.0, 0.0,],
        [0.0, 0.0, 1.0,],
        [0.0, 1.0, 0.0,],
        [0.0, 1.0, 1.0,],
        [1.0, 0.0, 0.0,],
        [1.0, 0.0, 1.0,],
        [1.0, 1.0, 0.0,],
        [1.0, 1.0, 1.0,],
])

pallet16 =  np.array([
        [0.0, 0.0, 0.0,], 
        [0.0, 1.0, 1.0,],
        [0.0, 0.0, 1.0,],
        [1.0, 0.0, 1.0,],
        [0.0, 0.5, 0.0,], 
        [0.5, 0.5, 0.5,],
        [0.0, 1.0, 0.0,],
        [0.5, 0.0, 0.0,],
        [0.0, 0.0, 0.5,],
        [0.5, 0.5, 0.0,],
        [0.5, 0.0, 0.5,],
        [1.0, 0.0, 0.0,],
        [0.75, 0.75, 0.75,],
        [0.0, 0.5, 0.5,],
        [1.0, 1.0, 1.0,], 
        [1.0, 1.0, 0.0,]
])


dx.p("Test funkcji colorFit dla palety 8 i 16 kolorów")
dx.p("Paleta kolorów")
dx.p("colorFit(np.array([0.25,0.25,0.5]),pallet8)")
dx.p(str(colorFit(np.array([0.25,0.25,0.5]),pallet8)))
dx.p("colorFit(np.array([0.25,0.25,0.5]),pallet16)")
dx.p(str(colorFit(np.array([0.25,0.25,0.5]),pallet16))),

dx.h("Wykorzystanie funkcji colorFit do kwantyzacji obrazu",1)

def kwant_colorFit(img,Pallete):
        out_img = img.copy()
        for w in range(img.shape[0]):
                for k in range(img.shape[1]):
                        out_img[w,k]=colorFit(img[w,k],Pallete)
        return out_img

images = ["GS_0001.tif","GS_0002.png","GS_0003.png"]


for image in images:
    ax, fig = plt.subplots(1, 4, figsize=(8, 5))
    img = plt.imread(image)
    if img[0].dtype == np.uint8:
        img = img.astype(np.float32) / 255.0
    
    ax.suptitle("Kwantyzacja z dopasowaniem do plalety")
    
    #  Original 0 0
    fig[0].imshow(img, cmap=plt.cm.gray)
    # Pallete 1-bit
    fig[1].set_title("Pallete 1-bit")
    fig[1].imshow(kwant_colorFit(img,pallet1), cmap=plt.cm.gray)
    # Pallete 2-bit
    fig[2].set_title("Pallete 2-bit")
    fig[2].imshow(kwant_colorFit(img,pallet2), cmap=plt.cm.gray)
    # Pallete 4-bit
    fig[3].set_title("Pallete 4-bit")
    fig[3].imshow(kwant_colorFit(img,pallet4), cmap=plt.cm.gray)
    dx.addImage(ax,5)
    
images = ["SMALL_0001.tif","SMALL_0009.jpg","SMALL_0007.jpg","SMALL_0004.jpg","SMALL_0006.jpg"]

for image in images:
    ax, fig = plt.subplots(1, 3, figsize=(8, 5))
    img = plt.imread(image)
    if img[0].dtype == np.uint8:
        img = img.astype(np.float32) / 255.0
    
    ax.suptitle("Kwantyzacja z dopasowaniem do plalety")
    
    #  Original 0 0
    fig[0].imshow(img, cmap=plt.cm.gray)
    # Pallete 8bit
    fig[1].set_title("Pallete 8-bit")
    fig[1].imshow(kwant_colorFit(img,pallet8), cmap=plt.cm.gray)
    # Pallete 16bit
    fig[2].set_title("Pallete 16-bit")
    fig[2].imshow(kwant_colorFit(img,pallet16), cmap=plt.cm.gray)
    dx.addImage(ax,5)
    
dx.h("Dithering",1)

def ditheringRandom(img):
    if len(img.shape) == 3:
        img2 = img[:,:,0].copy()
    else:
        img2 = img
    r = np.random.rand(img2.shape[0], img2.shape[1])
    o = (img2 >= r)*1
    return o

def ditheringOrdered(img,Pallete):
        n = 2
        output = img.copy()
        M = np.array([[0,8,2,10],
                    [12, 4, 14,6],
                    [3, 11, 1,9],
                    [15, 7, 13,5]])
        Mpre = (M+1) / (2*n)**2 - 0.5
        for w in range(output.shape[0]):
              for k in range(output.shape[1]):
                       output[w, k] = colorFit(img[w, k]+Mpre[w%(2*n),k%(2*n)],Pallete)
        return output

def ditheringFloyd(img,Pallete):
    output = img.copy()
    height = img.shape[0]
    width = img.shape[1]
    for w in range(height):
        for k in range(width):
            old_pixel = output[w, k].copy()
            new_pixel = colorFit(old_pixel, Pallete)
            output[w, k] = new_pixel
            quant_error = old_pixel - new_pixel
            if k + 1 < width:
                output[w, k + 1] = output[w, k + 1] + quant_error*7/16
            if (k - 1 >= 0) and (w + 1 < height):
                output[w + 1, k - 1] = output[w + 1, k - 1] + quant_error*3/16
            if w + 1 < height:
                output[w + 1, k] = output[w + 1, k] + quant_error*5/16
            if (w + 1 < height) and (k + 1 < width):
                output[w + 1, k + 1] = output[w + 1, k + 1] + quant_error*1/16
    return output




def calculateDithering(pallets, images):
    for image in images:
        for pallet in pallets:
            ax, fig = plt.subplots(1, 5, figsize=(10, 4))
            img = plt.imread(image)
            if img[0].dtype == np.uint8:
                img = img.astype(np.float32) / 255.0
            
            ax.suptitle("Dithering -"+pallet[0])
            
            #  Original 0 0
            fig[0].set_title("Original")
            fig[0].imshow(img, cmap=plt.cm.gray)
            # Pallete 1-bit
            fig[1].set_title("Kwantyzacja")
            fig[1].imshow(kwant_colorFit(img,pallet[1]), cmap=plt.cm.gray)
            # Random Dithering
            if pallet[0] == "Pallet 1-bit":
                fig[2].set_title("Losowy")
                fig[2].imshow(ditheringRandom(img), cmap=plt.cm.gray)
            else:
                fig[2].axis('off')
            # Ordered Dithering
            fig[3].set_title("Zorganizowany")
            fig[3].imshow(ditheringOrdered(img,pallet[1]), cmap=plt.cm.gray)
            # Floyd Dithering
            fig[4].set_title("Floyd")
            fig[4].imshow(ditheringFloyd(img,pallet[1]), cmap=plt.cm.gray)
            dx.addImage(ax,5)

pallets = [["Pallet 1-bit", pallet1], ["Pallet 2-bit", pallet2], ["Pallet 4-bit", pallet4]]
images = ["GS_0001.tif","GS_0003.png","GS_0002.png"]
calculateDithering(pallets, images)
pallets = [["Pallet 8-bit", pallet8], ["Pallet 16-bit", pallet16]]
images = ["SMALL_0001.tif","SMALL_0009.jpg","SMALL_0007.jpg","SMALL_0004.jpg","SMALL_0006.jpg"]
calculateDithering(pallets, images)

dx.save("Lab_3.docx")
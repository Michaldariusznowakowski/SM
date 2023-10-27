from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO

class DocxSave:
    def __init__(self):
        self.document = Document()
        
    def addHeading(self, text, level):
        self.document.add_heading(text, level)
    
    def addParagraph(self, text):
        self.document.add_paragraph(text)
    
    def addImage(self, image, width):
        if isinstance(image, plt.Figure):
            memfile = BytesIO()
            image.savefig(memfile)
            self.document.add_picture(memfile, width=Inches(width))
            memfile.close()
    
    def addCode(self, codePath):
        with open(codePath, 'r') as codeFile:
            codeText = codeFile.read()
            self.document.add_paragraph('Kod źródłowy:')
            codeParagraph = self.document.add_paragraph()
            codeRun = codeParagraph.add_run(codeText)
            codeRun.font.name = 'Courier New'
    
    def save(self, filename):
        self.document.save(filename)
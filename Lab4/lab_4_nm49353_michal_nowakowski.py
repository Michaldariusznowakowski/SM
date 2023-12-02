from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO
import sounddevice as sd
import soundfile as sf
from scipy.interpolate import interp1d
from scipy import signal
import scipy.fftpack
import os.path
####################################################################################################
# Lab 4
####################################################################################################

DEBUG = False # Brak zapisu do pliku, wyświetlanie wykresów
DISABLE_PLOTS = False # Wyłączenie wyświetlania wykresów
dataFolder = "data/"

for filename in os.listdir(dataFolder+"output/"):
    os.remove(dataFolder+"output/"+filename)    
    
####################################################################################################
class DocxSave:
    def __init__(self, debug=False):
        self.document = Document()
        self.debug = debug
        
    def h(self, text, level):
        if self.debug:
            print("h: "+text)
        else:
            self.document.add_heading(text, level)
    
    def p(self, text):
        if self.debug:
            print("p: "+text)
        else:
            self.document.add_paragraph(text)
    
    def addImage(self, image, width):
        if self.debug and DISABLE_PLOTS == False:
            plt.show()
        else:
            if isinstance(image, plt.Figure):
                memfile = BytesIO()
                image.savefig(memfile)
                self.document.add_picture(memfile, width=Inches(width))
                memfile.close()
                plt.close()
    
    def addCode(self, codePath):
        if not self.debug: 
            with open(codePath, 'r') as codeFile:
                codeText = codeFile.read()
                self.document.add_paragraph('Kod źródłowy:')
                codeParagraph = self.document.add_paragraph()
                codeRun = codeParagraph.add_run(codeText)
                codeRun.font.name = 'Courier New'
    
    def save(self, filename):
        if not self.debug:
            self.document.save(filename)

def plotAudio(Signal, Fs, TimeMargin=[0, 0.04]):
    start = int(TimeMargin[0] * Fs)
    end = int(TimeMargin[1] * Fs)
    signalFragment = Signal[start:end]
    # time in sec
    time = np.linspace(0, len(signalFragment) / Fs, num=len(signalFragment))
    fig, axs = plt.subplots(2, 1, figsize=(10, 15))
    axs[0].plot(time, signalFragment)
    axs[0].set_title('Sygnał dźwiękowy')
    axs[0].set_xlabel('Czas [s]')
    axs[0].set_ylabel('Amplituda')
    spectrum = 20 * np.log10(np.abs(np.fft.fft(signalFragment))+np.finfo(np.float32).eps)
    freq = np.fft.fftfreq(len(signalFragment), 1 / Fs)
    halfSpectrum = spectrum[:len(spectrum)//2]
    halfFreq = freq[:len(freq)//2]
    axs[1].plot(halfFreq, halfSpectrum)
    axs[1].set_title('Widmo dźwiękowe (połowa)')
    axs[1].set_xlabel('Częstotliwość [Hz]')
    axs[1].set_ylabel('Amplituda [dB]')
    fig.tight_layout()
    return fig

def saveAudio(Signal, Fs, filename):
    file = dataFolder+"output/"+filename.replace(".wav", "")+".wav"
    if os.path.isfile(file):
        os.remove(file)
                
    sf.write(file, Signal, Fs)

dx = DocxSave(debug=DEBUG)
####################################################################################################

def Kwant(data,bit):
    output = data.astype(np.float32)
    beg = np.iinfo(data.dtype).min
    end   = np.iinfo(data.dtype).max
    d=2**bit-1
    if np.issubdtype(data.dtype,np.integer):
        output = (output-beg)/(end-beg)
        output = np.round(output*d)/d
        output = ((output*(end-beg))+beg).astype(data.dtype)
    else :
        output = np.round(output*d)/d
    return output
        
        
def Decymacja(data,fs,decymacja):
    output = data[::decymacja]
    newFs = fs//decymacja
    return output, newFs

def Interpolacja(data,fs,fs2,typ):
    duration = data.shape[0]/fs
    x_old = np.arange(0, duration, 1/fs)
    x_new = np.arange(0, duration, 1/fs2)
    interp_func = interp1d(x_old, data, typ, fill_value="extrapolate")
    return interp_func(x_new).astype(data.dtype),fs2

dx.h("Wnioski i obserwacje",0)  
dx.h("Zadanie 2 Część 1",2)
dx.h("Kwantyzacja",3)
dx.p("Zwiększenie ilości bitów, wpływa pozytywnie na jakość dźwięku")
dx.p("Dla 4 bitów dźwięk jest zniekształcony, a dla 8 bitów występuje wyraźna poprawa, ale na wykresie widma widać, że dźwięk jest zniekształcony")
dx.p("Dla 16 i 24 bitów nie ma wyraźnej różnicy w jakości dźwięku")
dx.h("Decymacja",3)
dx.p("Wraz ze zwiększeniem ilości redukowanych próbek, efektywnie zmniejszamy częstotliwość próbkowania, co powoduje zniekształcenia dźwięku")
dx.p("Dodatkowo zmniejszamy rozmiar pliku, co jest pożądane")
dx.h("Interpolacja liniowa i nieliniowa",3)
dx.p("Wraz ze zwiększeniem ilości interpolowanych próbek, wykres dźwięku staje się bardziej gładki, a widmo dźwięku jest mniej zniekształcone")
dx.h("Zadanie 2 Część 2",2  )
dx.h("Dla pliku sing_high1.wav",3)
dx.h("Kwantyzacja",4)
dx.p("Dla 4 bitów, dźwięk jest zniekształcony i o wiele głośniejszy niż oryginał")
dx.p("Dla 8 bitów, dźwięk jest zniekształcony, ale nie tak bardzo jak dla 4 bitów")

dx.h("Decymacja",4)
dx.p("Decymacji z krokiem 4, dźwięk jest niższy i nieco zniekształcony")
dx.p("Dla decymacji z krokiem 6, dźwięk jest bardziej zniekształcony niż dla kroku 4")
dx.p("Dla decymacji z krokiem 10, dźwięk jest jeszcze bardziej zniekształcony niż dla kroku 6")
dx.p("Dla decymacji z krokiem 24, ciężko usłyszeć originalny dźwięk, w tle słychać ufo ;)")

dx.h("Interpolacja liniowa i nieliniowa",4)
dx.p("Dla interpolacji na 4000 próbek, dźwięk jest zniekształcony, wytłumiony i nieco niższy")
dx.p("Dla interpolacji na 8000 próbek, słychać wyraźną poprawę, ale dźwięk jest wytłumiony")
dx.p("Dla interpolacji na 11999, 16000, 16953 próbek, dźwięk jest coraz bardziej podobny do oryginału")

dx.h("Dla pliku sing_low1.wav",3)
dx.h("Kwantyzacja",4)
dx.p("Dla 4 bitów, dźwięk jest zniekształcony i o WIELE głośniejszy niż oryginał")
dx.p("Dla 8 bitów, dźwięk jest zniekształcony, ale nie tak bardzo jak dla 4 bitów! Jest o wiele cichszy!")

dx.h("Decymacja",4)
dx.p("Dla decymacji z krokiem 4, dźwięk jest niższy")
dx.p("Dla decymacji z krokiem 6, dźwięk jest niższy i nieco zniekształcony")
dx.p("Dla decymacji z krokiem 10, dźwięk jest bardziej zniekształcony niż dla kroku 6 już w tle słuchać odgłosy ufo")
dx.p("Dla decymacji z krokiem 24, ciężko usłyszeć originalny dźwięk)")

dx.h("Interpolacja liniowa i nieliniowa",4)
dx.p("Dla interpolacji na 4000 próbek, bardzo niska jakość dźwięku, dźwięk jest mocno wytłumiony")
dx.p("Dla interpolacji na 8000 próbek, o wiele lepsza jakość dźwięku, bliska oryginałowi")
dx.p("Dla interpolacji na 11999, 16000, 16953 próbek, dźwięk jest coraz bardziej podobny do oryginału")
dx.h("Dla pliku sing_medium1.wav",3)
dx.h("Kwantyzacja",4)
dx.p("Dla 4 bitów, ostry przester, duże szumy, i niska jakość dźwięku")
dx.p("Dla 8 bitów, duża poprawa jakości dźwięku, ale dźwięk jest wytłumiony")
dx.h("Decymacja",4)
dx.p("Dla decymacji z krokiem 4, dźwięk jest niższy i nieco zniekształcony")
dx.p("Dla decymacji z krokiem 6, dźwięk jest niższy ")
dx.h("Interpolacja liniowa i nieliniowa",4)
dx.p("Dla interpolacji na 4000 próbek, bardzo niska jakość dźwięku, mocne wytłumienie dźwięku")
dx.p("Dla interpolacji na 8000 próbek, wyraźna poprawa jakości dźwięku")
dx.p("Dla interpolacji na 11999, 16000, 16953 próbek, dźwięk jest coraz bardziej podobny do oryginału")

dx.h("Wnioski",3)
dx.h("Kwantyzacja",4)
dx.p("Wraz ze zwiększeniem ilości bitów, dźwięk jest coraz mniej zniekształcony i coraz bardziej podobny do oryginału")
dx.h("Decymacja",4)
dx.p("Wraz ze zwiększeniem ilości decymowanych próbek, dźwięk jest coraz bardziej zniekształcony")
dx.h("Interpolacja liniowa i nieliniowa",4)
dx.p("Wraz ze zwiększeniem ilości interpolowanych próbek, dźwięk jest coraz mniej zniekształcony i coraz bardziej podobny do oryginału")
dx.h("Zadanie 2",0)


###############################################
# Część 1
dx.h("Zadanie   2 Część 1",0)
sygnalBitowyLista=[4,8,16,24]
decymacjaLista=[4,8,16,24]
interpolacjaLista=[2000,4000,8000,11999,1600,16953,24000,24000,41000]



sinFiles=["sin_60Hz.wav","sin_440Hz.wav","sin_8000Hz.wav","sin_combined.wav"]

for filename in sinFiles:
    file = dataFolder+filename
    data, fs = sf.read(file, dtype=np.int32)
    dx.h(f'Analiza pliku {filename}', level=1)
    dx.p('Przed modyfikacją')
    fig = plotAudio(data, fs)
    dx.addImage(fig, width=5)
    dx.p('Po kwantyzacji')
    for bit in sygnalBitowyLista:
        dx.p(f'Kwantyzacja na {bit} bitów')
        d = Kwant(data,bit)
        fig = plotAudio(d, fs)
        dx.addImage(fig, width=5)
    dx.p('Po decymacji')
    for decymacja in decymacjaLista:
        dx.p(f'Decymacja na {decymacja}')
        d, nfs = Decymacja(data,fs,decymacja)
        fig = plotAudio(d, nfs)
        dx.addImage(fig, width=5)
    dx.p('Po interpolacji liniowej')
    for interpolacja in interpolacjaLista:
        dx.p(f'Interpolacja na {interpolacja}')
        d, nfs = Interpolacja(data,fs,interpolacja,'linear')
        fig = plotAudio(d, nfs)
        dx.addImage(fig, width=5)
    dx.p('Po interpolacji nielinowej')
    for interpolacja in interpolacjaLista:
        dx.p(f'Interpolacja na {interpolacja}')
        d,nfs = Interpolacja(data,fs,interpolacja,'quadratic')
        fig = plotAudio(d, nfs)
        dx.addImage(fig, width=5)
        
        

# ################################################
# # Część 2
dx.h("Zadanie 2 Część 2",0)
dx.p("Pliki wynikowe dodane osobno obok sprawozdania")

sygnalBitowyLista=[4,8]
decymacjaLista=[4,6,10,24]
interpolacjaLista=[4000,8000,11999,16000,16953]

files = ["sing_low1.wav","sing_high1.wav","sing_medium1.wav"]
for file in files:
    data, fs = sf.read(dataFolder+file, dtype=np.int32)
    dx.h(f'Analiza pliku {file}', level=1)
    dx.p('Przed modyfikacją')
    fig = plotAudio(data, fs)
    dx.addImage(fig, width=5)
    dx.p('Po kwantyzacji')
    for bit in sygnalBitowyLista:
        dx.p(f'Kwantyzacja na {bit} bitów')
        d = Kwant(data,bit)
        fig = plotAudio(d, fs)
        dx.addImage(fig, width=5)
        saveAudio(d, fs, file+"_kwant_"+str(bit))
    dx.p('Po decymacji')
    for decymacja in decymacjaLista:
        dx.p(f'Decymacja na {decymacja}')
        d, nfs = Decymacja(data,fs,decymacja)
        fig = plotAudio(d, nfs)
        dx.addImage(fig, width=5)
        saveAudio(d, nfs, file+"_decymacja_"+str(decymacja))
    dx.p('Po interpolacji liniowej')
    for interpolacja in interpolacjaLista:
        dx.p(f'Interpolacja na {interpolacja}')
        d, nfs = Interpolacja(data,fs,interpolacja,'linear')
        fig = plotAudio(d, nfs)
        dx.addImage(fig, width=5)
        saveAudio(d, nfs, file+"_interpolacja_"+str(interpolacja))
    dx.p('Po interpolacji nielinowej')
    for interpolacja in interpolacjaLista:
        dx.p(f'Interpolacja na {interpolacja}')
        d,nfs = Interpolacja(data,fs,interpolacja,'quadratic')
        fig = plotAudio(d, nfs)
        dx.addImage(fig, width=5)
        saveAudio(d, nfs, file+"_interpolacja_nielin_"+str(interpolacja))

if os.path.isfile(dataFolder+"output/Lab_4.docx"):
    os.remove(dataFolder+"output/Lab_4.docx")
dx.save(dataFolder+"output/Lab_4.docx")
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO
import sounddevice as sd
import soundfile as sf
from scipy.interpolate import interp1d
from scipy import signal
import scipy.fftpack

class DocxSave:
    def __init__(self, debug=False):
        self.document = Document()
        self.debug = debug
        
    def h(self, text, level):
        if self.debug:
            print("h: "+text)
        else:
            self.document.add_heading(text, level)
    
    def p(self, text):
        if self.debug:
            print("p: "+text)
        else:
            self.document.add_paragraph(text)
    
    def addImage(self, image, width):
        if self.debug:
            plt.show()
        else:
            if isinstance(image, plt.Figure):
                memfile = BytesIO()
                image.savefig(memfile)
                self.document.add_picture(memfile, width=Inches(width))
                memfile.close()
                plt.close()
    
    def addCode(self, codePath):
        if not self.debug: 
            with open(codePath, 'r') as codeFile:
                codeText = codeFile.read()
                self.document.add_paragraph('Kod źródłowy:')
                codeParagraph = self.document.add_paragraph()
                codeRun = codeParagraph.add_run(codeText)
                codeRun.font.name = 'Courier New'
    
    def save(self, filename):
        if not self.debug:
            self.document.save(filename)


####################################################################################################
# Lab 4
####################################################################################################

DEBUG = False # Brak zapisu do pliku, wyświetlanie wykresów
dx = DocxSave(debug=DEBUG)
images_folder = "data/"

####################################################################################################

def Kwant(data,bit):
    output = data.astype(np.float32)
    beg = np.iinfo(data.dtype).min
    end   = np.iinfo(data.dtype).max
    d=2**bit-1
    if np.issubdtype(data.dtype,np.integer):
        output = (output-beg)/(end-beg)
        output = np.round(output*d)/d
        output = ((output*(end-beg))+beg).astype(data.dtype)
    else :
        output = np.round(output*d)/d
        
def decymacja(data,fs,fs2):
    return data[::fs/fs2]

def interpolacja(data,fs,fs2,typ): # funkcja z biblioteki scipy
    duration = data.shape[0]/fs
    x_old = np.arange(0, duration, 1/fs)
    x_new = np.arange(0, duration, 1/fs2)
    interp_func = interp1d(x_old, data, typ, fill_value="extrapolate")
    return interp_func(x_new).astype(data.dtype)


fig = plt.figure()
fig.tight_layout()
dx.addImage(fig, width=3)

dx.addParagraph("Widmo")

data, fs = sf.read('Lab1/data/sin_440Hz.wav', dtype=np.int32)

fig, axs = plt.subplots(2, 1, figsize=(3, 4))

axs[0].plot(np.arange(0,data.shape[0])/fs,data)

yf = scipy.fftpack.fft(data)
axs[1].plot(np.arange(0,fs,1.0*fs/(yf.size)),np.abs(yf))
dx.addImage(fig, width=3)

dx.addParagraph("Moduł widma, zmieniony rozmiar transformaty Fouriera")

fsize=2**8
fig, axs = plt.subplots(2, 1, figsize=(3, 4))

axs[0].plot(np.arange(0,data.shape[0])/fs,data)

yf = scipy.fftpack.fft(data,fsize)
axs[1].plot(np.arange(0,fs,fs/fsize),np.abs(yf))
dx.addImage(fig, width=3)

dx.addParagraph("Wyświetlenie modułu widma, wzięcie tylko połowy")

fig, axs = plt.subplots(2, 1, figsize=(3, 4))

axs[0].plot(np.arange(0,data.shape[0])/fs,data)
yf = scipy.fftpack.fft(data,fsize)
axs[1].plot(np.arange(0,fs/2,fs/fsize),np.abs(yf[:fsize//2]))
dx.addImage(fig, width=3)


dx.addParagraph("Przeskalowanie wartości widma do skali decybelowej (dB)")

fig, axs = plt.subplots(2, 1, figsize=(3, 4))


axs[0].plot(np.arange(0,data.shape[0])/fs,data)
yf = scipy.fftpack.fft(data,fsize)
axs[1].plot(np.arange(0,fs/2,fs/fsize),20*np.log10( np.abs(yf[:fsize//2])))

dx.addImage(fig, width=3)

dx.addHeading("Zadanie 2",0)

def plotAudio(Signal, Fs, TimeMargin=[0, 0.02]):
    start = int(TimeMargin[0] * Fs)
    end = int(TimeMargin[1] * Fs)
    signalFragment = Signal[start:end]
    # time in sec
    time = np.arange(0, len(signalFragment) / Fs, 1 / Fs)
    fig, axs = plt.subplots(2, 1, figsize=(3, 4))
    axs[0].plot(time, signalFragment)
    axs[0].set_title('Sygnał dźwiękowy')
    axs[0].set_xlabel('Czas [s]')
    axs[0].set_ylabel('Amplituda')
    # widmo w dB
    spectrum = 20 * np.log10(np.abs(np.fft.fft(signalFragment)) / len(signalFragment))
    freq = np.fft.fftfreq(len(signalFragment), 1 / Fs)
    # połowa widma
    halfSpectrum = spectrum[:len(spectrum)//2]
    halfFreq = freq[:len(freq)//2]
    axs[1].plot(halfFreq, halfSpectrum)
    axs[1].set_title('Widmo dźwiękowe (połowa)')
    axs[1].set_xlabel('Częstotliwość [Hz]')
    axs[1].set_ylabel('Amplituda [dB]')
    fig.tight_layout()
    return fig




################################################
sygnalBitowyLista=[4,8,16,24]
decymacjaLista=[4,8,16,24]
interpolacjaLista=[2000,4000,8000,11999,1600,16953,24000,24000,41000]

################################################
# Część 2
sygnalBitowyLista=[4,8]
decymacjaLista=[4,6,10,24]
interpolacjaLista=[4000,8000,11999,16000,16953]

Signal, Fs = sf.read('Lab1/data/sin_440Hz.wav')
fig = plotAudio(Signal, Fs)

dx.addImage(fig, width=3)
dx.addHeading("Zadanie 3",0)
files=["sin_60Hz.wav",
"sin_440Hz.wav",
"sin_8000Hz.wav",
"sin_combined.wav"]
fsizeValues = [2**8, 2**12, 2**16]
for filename in files:
    for fsize in fsizeValues:
            data, fs = sf.read("Lab1/data/"+filename, dtype=np.int32)
            # obliczanie widma dla danego fsize
            yf = np.fft.fft(data, fsize)
            freqs = np.fft.fftfreq(fsize, 1.0 / fs)
            
            maxAmplitudeIndex = np.argmax(np.abs(yf))
            maxAmplitude = np.abs(yf[maxAmplitudeIndex])
            maxFrequency = freqs[maxAmplitudeIndex]
            
            fig = plotAudio(Signal, Fs)
            dx.addImage(fig, width=3)
            dx.addHeading(f'Analiza dla fsize={fsize}', level=1)
            dx.addParagraph(f'Najwyższa amplituda: {maxAmplitude}')
            dx.addParagraph(f'Dla częstotliwości: {maxFrequency} Hz')
